"""
File Parser Utilities
Extract text content from various file types for AI processing
"""
import os
import csv
import io
from typing import Dict, Optional
import PyPDF2
import docx
import json


def extract_file_content(filepath: str, content_type: Optional[str] = None) -> Dict[str, any]:
    """
    Extract text content from various file types
    
    Args:
        filepath: Path to the file
        content_type: MIME type of the file
    
    Returns:
        Dict with success status, content, and metadata
    """
    try:
        ext = os.path.splitext(filepath)[1].lower()
        
        # Text files
        if ext in ['.txt', '.md', '.py', '.js', '.json', '.xml', '.html', '.css']:
            return extract_text_file(filepath)
        
        # CSV files
        elif ext == '.csv':
            return extract_csv_file(filepath)
        
        # PDF files
        elif ext == '.pdf':
            return extract_pdf_file(filepath)
        
        # Word documents
        elif ext in ['.docx', '.doc']:
            return extract_docx_file(filepath)
        
        # JSON files
        elif ext == '.json':
            return extract_json_file(filepath)
        
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {ext}",
                "content": None
            }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "content": None
        }


def extract_text_file(filepath: str) -> Dict:
    """Extract content from plain text files"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return {
            "success": True,
            "content": content,
            "content_type": "text",
            "lines": len(content.split('\n')),
            "chars": len(content)
        }
    except UnicodeDecodeError:
        # Try with different encoding
        with open(filepath, 'r', encoding='latin-1') as f:
            content = f.read()
        
        return {
            "success": True,
            "content": content,
            "content_type": "text",
            "lines": len(content.split('\n')),
            "chars": len(content)
        }


def extract_csv_file(filepath: str) -> Dict:
    """Extract and format content from CSV files"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            csv_reader = csv.reader(f)
            rows = list(csv_reader)
        
        if not rows:
            return {
                "success": False,
                "error": "CSV file is empty",
                "content": None
            }
        
        # Format as table for AI
        headers = rows[0]
        data_rows = rows[1:]
        
        # Create formatted text representation
        content = f"CSV File Content:\n\n"
        content += f"Headers: {', '.join(headers)}\n"
        content += f"Total rows: {len(data_rows)}\n\n"
        
        # Add first 50 rows (to avoid token limits)
        content += "Data:\n"
        for i, row in enumerate(data_rows[:50]):
            content += f"Row {i+1}: {', '.join(str(cell) for cell in row)}\n"
        
        if len(data_rows) > 50:
            content += f"\n... and {len(data_rows) - 50} more rows"
        
        return {
            "success": True,
            "content": content,
            "content_type": "csv",
            "rows": len(data_rows),
            "columns": len(headers),
            "headers": headers
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "content": None
        }


def extract_pdf_file(filepath: str) -> Dict:
    """Extract text content from PDF files"""
    try:
        with open(filepath, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            
            content = f"PDF File Content ({num_pages} pages):\n\n"
            
            # Extract text from each page (limit to first 20 pages)
            for i, page in enumerate(pdf_reader.pages[:20]):
                text = page.extract_text()
                content += f"--- Page {i+1} ---\n{text}\n\n"
            
            if num_pages > 20:
                content += f"\n... and {num_pages - 20} more pages"
        
        return {
            "success": True,
            "content": content,
            "content_type": "pdf",
            "pages": num_pages,
            "chars": len(content)
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "content": None
        }


def extract_docx_file(filepath: str) -> Dict:
    """Extract text content from Word documents"""
    try:
        doc = docx.Document(filepath)
        
        content = "Word Document Content:\n\n"
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                content += para.text + "\n"
        
        return {
            "success": True,
            "content": content,
            "content_type": "docx",
            "paragraphs": len(doc.paragraphs),
            "chars": len(content)
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "content": None
        }


def extract_json_file(filepath: str) -> Dict:
    """Extract and format JSON content"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Pretty print JSON
        content = f"JSON File Content:\n\n{json.dumps(data, indent=2)}"
        
        return {
            "success": True,
            "content": content,
            "content_type": "json",
            "chars": len(content)
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "content": None
        }


def summarize_file_content(content: str, max_tokens: int = 3000) -> str:
    """
    Truncate file content if too long
    ~4 characters per token
    """
    max_chars = max_tokens * 4
    
    if len(content) <= max_chars:
        return content
    
    # Truncate and add notice
    truncated = content[:max_chars]
    return truncated + f"\n\n[Content truncated - showing first {max_tokens} tokens of file]"
